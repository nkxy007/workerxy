"""
Document processing module - handles PDF, Markdown, and text files
"""
import re
from typing import List, Dict, Any, Optional
from pathlib import Path
from dataclasses import dataclass
from enum import Enum

import PyPDF2
from docx import Document as DocxDocument


class DocumentType(Enum):
    """Supported document types"""
    PDF = "pdf"
    DOCX = "docx"
    MARKDOWN = "md"
    TEXT = "txt"
    CODE = "code"


@dataclass
class DocumentPage:
    """Represents a single page or section"""
    content: str
    page_num: Optional[int] = None
    metadata: Dict[str, Any] = None
    
    def __post_init__(self):
        if self.metadata is None:
            self.metadata = {}


@dataclass
class Document:
    """Processed document with pages/sections"""
    doc_id: str
    doc_type: DocumentType
    pages: List[DocumentPage]
    total_pages: int
    file_path: str
    metadata: Dict[str, Any] = None
    
    def __post_init__(self):
        if self.metadata is None:
            self.metadata = {}


class DocumentProcessor:
    """Process various document types"""
    
    @staticmethod
    def detect_type(file_path: str) -> DocumentType:
        """Detect document type from file extension"""
        suffix = Path(file_path).suffix.lower()
        
        type_map = {
            '.pdf': DocumentType.PDF,
            '.docx': DocumentType.DOCX,
            '.doc': DocumentType.DOCX,
            '.md': DocumentType.MARKDOWN,
            '.markdown': DocumentType.MARKDOWN,
            '.txt': DocumentType.TEXT,
            '.py': DocumentType.CODE,
            '.js': DocumentType.CODE,
            '.java': DocumentType.CODE,
            '.cpp': DocumentType.CODE,
            '.c': DocumentType.CODE,
        }
        
        return type_map.get(suffix, DocumentType.TEXT)
    
    @staticmethod
    def process_pdf(file_path: str) -> Document:
        """Process PDF document"""
        pages = []
        
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            total_pages = len(pdf_reader.pages)
            
            for page_num, page in enumerate(pdf_reader.pages, start=1):
                text = page.extract_text()
                
                pages.append(DocumentPage(
                    content=text,
                    page_num=page_num,
                    metadata={'doc_type': 'pdf', 'is_paginated': True}
                ))
        
        doc_id = Path(file_path).stem
        return Document(
            doc_id=doc_id,
            doc_type=DocumentType.PDF,
            pages=pages,
            total_pages=total_pages,
            file_path=file_path,
            metadata={'is_paginated': True}
        )
    
    @staticmethod
    def process_docx(file_path: str) -> Document:
        """Process DOCX document"""
        doc = DocxDocument(file_path)
        pages = []
        
        # DOCX doesn't have natural pages, treat each paragraph as potential chunk
        current_page = []
        page_num = 1
        chars_per_page = 2000  # Approximate characters per "page"
        current_chars = 0
        
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            
            current_page.append(text)
            current_chars += len(text)
            
            # Create a new "page" when we hit the character limit
            if current_chars >= chars_per_page:
                pages.append(DocumentPage(
                    content='\n'.join(current_page),
                    page_num=page_num,
                    metadata={'doc_type': 'docx', 'is_paginated': True}
                ))
                current_page = []
                current_chars = 0
                page_num += 1
        
        # Add remaining content
        if current_page:
            pages.append(DocumentPage(
                content='\n'.join(current_page),
                page_num=page_num,
                metadata={'doc_type': 'docx', 'is_paginated': True}
            ))
        
        doc_id = Path(file_path).stem
        return Document(
            doc_id=doc_id,
            doc_type=DocumentType.DOCX,
            pages=pages,
            total_pages=len(pages),
            file_path=file_path,
            metadata={'is_paginated': True}
        )
    
    @staticmethod
    def process_markdown(file_path: str) -> Document:
        """Process Markdown with hierarchical semantic chunking"""
        with open(file_path, 'r', encoding='utf-8') as f:
            content = f.read()
        
        # Parse headers to create hierarchical structure
        sections = DocumentProcessor._parse_markdown_sections(content)
        
        doc_id = Path(file_path).stem
        return Document(
            doc_id=doc_id,
            doc_type=DocumentType.MARKDOWN,
            pages=sections,
            total_pages=len(sections),
            file_path=file_path,
            metadata={'is_paginated': False, 'hierarchical': True}
        )
    
    @staticmethod
    def _parse_markdown_sections(content: str) -> List[DocumentPage]:
        """Parse markdown into hierarchical sections based on headers"""
        sections = []
        
        # Split by headers (H1-H6)
        header_pattern = r'^(#{1,6})\s+(.+?)$'
        lines = content.split('\n')
        
        current_section = []
        current_header = None
        current_level = 0
        section_id = 0
        
        for line in lines:
            match = re.match(header_pattern, line, re.MULTILINE)
            
            if match:
                # Save previous section
                if current_section:
                    sections.append(DocumentPage(
                        content='\n'.join(current_section),
                        page_num=None,
                        metadata={
                            'section_id': section_id,
                            'section_title': current_header,
                            'section_level': current_level,
                            'doc_type': 'markdown',
                            'is_paginated': False
                        }
                    ))
                    section_id += 1
                
                # Start new section
                level = len(match.group(1))
                header = match.group(2)
                current_section = [line]
                current_header = header
                current_level = level
            else:
                current_section.append(line)
        
        # Add last section
        if current_section:
            sections.append(DocumentPage(
                content='\n'.join(current_section),
                page_num=None,
                metadata={
                    'section_id': section_id,
                    'section_title': current_header,
                    'section_level': current_level,
                    'doc_type': 'markdown',
                    'is_paginated': False
                }
            ))
        
        return sections if sections else [DocumentPage(
            content=content,
            metadata={'section_id': 0, 'doc_type': 'markdown', 'is_paginated': False}
        )]
    
    @staticmethod
    def process_text(file_path: str) -> Document:
        """Process plain text with semantic chunking"""
        with open(file_path, 'r', encoding='utf-8') as f:
            content = f.read()
        
        # Split by double newlines (paragraph boundaries)
        sections = DocumentProcessor._parse_text_sections(content)
        
        doc_id = Path(file_path).stem
        return Document(
            doc_id=doc_id,
            doc_type=DocumentType.TEXT,
            pages=sections,
            total_pages=len(sections),
            file_path=file_path,
            metadata={'is_paginated': False}
        )
    
    @staticmethod
    def _parse_text_sections(content: str, chars_per_section: int = 1500) -> List[DocumentPage]:
        """Parse text into sections based on paragraph boundaries"""
        sections = []
        
        # Split by double newlines
        paragraphs = re.split(r'\n\s*\n', content)
        
        current_section = []
        current_chars = 0
        section_id = 0
        
        for para in paragraphs:
            para = para.strip()
            if not para:
                continue
            
            para_chars = len(para)
            
            # If adding this paragraph exceeds limit, save current section
            if current_chars + para_chars > chars_per_section and current_section:
                sections.append(DocumentPage(
                    content='\n\n'.join(current_section),
                    page_num=None,
                    metadata={
                        'section_id': section_id,
                        'doc_type': 'text',
                        'is_paginated': False
                    }
                ))
                current_section = []
                current_chars = 0
                section_id += 1
            
            current_section.append(para)
            current_chars += para_chars
        
        # Add remaining content
        if current_section:
            sections.append(DocumentPage(
                content='\n\n'.join(current_section),
                page_num=None,
                metadata={
                    'section_id': section_id,
                    'doc_type': 'text',
                    'is_paginated': False
                }
            ))
        
        return sections if sections else [DocumentPage(
            content=content,
            metadata={'section_id': 0, 'doc_type': 'text', 'is_paginated': False}
        )]
    
    @classmethod
    def process_document(cls, file_path: str) -> Document:
        """
        Process any supported document type
        
        Args:
            file_path: Path to document file
            
        Returns:
            Processed Document object
        """
        doc_type = cls.detect_type(file_path)
        
        processors = {
            DocumentType.PDF: cls.process_pdf,
            DocumentType.DOCX: cls.process_docx,
            DocumentType.MARKDOWN: cls.process_markdown,
            DocumentType.TEXT: cls.process_text,
            DocumentType.CODE: cls.process_text,  # Use text processor for code
        }
        
        processor = processors.get(doc_type)
        if not processor:
            raise ValueError(f"Unsupported document type: {doc_type}")
        
        return processor(file_path)
